import pytest
import tempfile
import os
import anonymizer as anonymizer_module
from anonymizer import Anonymizer


class TestAnonymizer:
    def setup_method(self):
        self.config = {
            "auto_detect": True,
            "sensitivity": "ask_uncertain",
            "custom_terms": {
                "schools": ["國立OO大學"],
                "departments": ["資訊工程學系"],
            },
            "substring_match": True,
            "scan_paths": ["/tmp/test_scan/"],
            "file_types": [".md", ".txt", ".docx"],
            "persist_mapping": False,
            "max_file_pages": 50,
        }

    def test_anonymize_text_with_custom_terms(self):
        anon = Anonymizer(config=self.config, session_id="test1", use_ner=False)
        text = "國立OO大學資訊工程學系的報告"
        result, summary = anon.anonymize_text(text)
        assert "__ANON:SCHOOL_001__" in result
        assert "__ANON:DEPT_001__" in result
        assert "國立OO大學" not in result

    def test_anonymize_text_with_regex(self):
        anon = Anonymizer(config=self.config, session_id="test2", use_ner=False)
        text = "聯絡 test@school.edu.tw 或撥打 0912-345-678"
        result, summary = anon.anonymize_text(text)
        assert "__ANON:EMAIL_001__" in result
        assert "__ANON:PHONE_001__" in result
        assert "test@school.edu.tw" not in result

    def test_anonymize_preserves_structure(self):
        anon = Anonymizer(config=self.config, session_id="test3", use_ner=False)
        text = "# 標題\n\n國立OO大學的研究報告\n\n## 作者\n\ntest@test.com"
        result, summary = anon.anonymize_text(text)
        assert result.startswith("# 標題")
        assert "## 作者" in result

    def test_anonymize_file(self):
        anon = Anonymizer(config=self.config, session_id="test4", use_ner=False)
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write("國立OO大學 test@test.com")
            path = f.name
        try:
            anon_path, summary = anon.anonymize_file(path)
            assert os.path.exists(anon_path)
            with open(anon_path, 'r') as af:
                content = af.read()
            assert "國立OO大學" not in content
            assert "__ANON:SCHOOL_001__" in content
        finally:
            os.unlink(path)
            if anon_path and os.path.exists(anon_path):
                os.unlink(anon_path)

    def test_anonymize_docx_file_preserves_format(self):
        from docx import Document

        anon = Anonymizer(config=self.config, session_id="test_docx", use_ner=False)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name

        anon_path = None
        try:
            doc = Document()
            doc.add_paragraph("國立OO大學 test@test.com")
            doc.save(path)

            anon_path, summary = anon.anonymize_file(path)
            out_doc = Document(anon_path)
            text = "\n".join(p.text for p in out_doc.paragraphs)

            assert "國立OO大學" not in text
            assert "__ANON:SCHOOL_001__" in text
            assert anon_path.endswith(".docx")
        finally:
            os.unlink(path)
            if anon_path and os.path.exists(anon_path):
                os.unlink(anon_path)

    def test_no_pii_returns_unchanged(self):
        anon = Anonymizer(config=self.config, session_id="test5", use_ner=False)
        text = "今天天氣很好，出去走走"
        result, summary = anon.anonymize_text(text)
        assert result == text
        assert "未發現" in summary

    def test_right_to_left_replacement(self):
        anon = Anonymizer(config=self.config, session_id="test6", use_ner=False)
        text = "國立OO大學和資訊工程學系"
        result, summary = anon.anonymize_text(text)
        assert "__ANON:SCHOOL_001__" in result
        assert "__ANON:DEPT_001__" in result
        assert "和" in result

    def test_anonymize_file_to_text_temp_respects_max_file_pages(self, monkeypatch):
        class FakePdfParser:
            def __init__(self):
                self.calls = []

            def parse(self, file_path, max_pages=50):
                self.calls.append(max_pages)
                return "電話 0912345678"

        parser = FakePdfParser()
        # Patch get_parser in the actual globals dict used by Anonymizer methods.
        # When pytest imports anonymizer as a package (__init__.py), importlib
        # loads anonymizer.py into a separate _core module. Anonymizer's method
        # globals point to _core, not the package namespace, so we must patch
        # the method's own globals to intercept the call.
        target = Anonymizer.anonymize_file_to_text_temp.__globals__
        monkeypatch.setitem(target, "get_parser", lambda _path: parser)

        config = dict(self.config)
        config["max_file_pages"] = 3
        anon = Anonymizer(config=config, session_id="test_pdf_limit", use_ner=False)

        anon_path, _summary = anon.anonymize_file_to_text_temp("/tmp/example.pdf")
        try:
            assert parser.calls == [3]
            assert anon_path is not None
        finally:
            if anon_path and os.path.exists(anon_path):
                os.unlink(anon_path)
