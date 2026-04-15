"""Tests for batch.py — folder-level batch processing."""

import os

import pytest

from batch import BatchResult, _collect_files, run_batch


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _create_test_tree(tmp_path):
    """Create a test directory tree with sample files."""
    # Text files
    (tmp_path / "doc1.txt").write_text("張三的電話是0912345678", encoding="utf-8")
    (tmp_path / "doc2.md").write_text("無敏感資料", encoding="utf-8")

    # Subfolder
    sub = tmp_path / "sub"
    sub.mkdir()
    (sub / "doc3.txt").write_text("李四 0987654321", encoding="utf-8")

    # Unsupported file
    (tmp_path / "readme.log").write_text("some log data", encoding="utf-8")

    return tmp_path


# ---------------------------------------------------------------------------
# File collection
# ---------------------------------------------------------------------------

class TestCollectFiles:
    def test_collects_matching_files(self, tmp_path):
        _create_test_tree(tmp_path)
        files = _collect_files(str(tmp_path), [".txt", ".md"])
        basenames = [os.path.basename(f) for f in files]
        assert "doc1.txt" in basenames
        assert "doc2.md" in basenames
        assert "doc3.txt" in basenames
        assert "readme.log" not in basenames

    def test_empty_dir(self, tmp_path):
        files = _collect_files(str(tmp_path), [".txt"])
        assert files == []

    def test_no_matching_types(self, tmp_path):
        _create_test_tree(tmp_path)
        files = _collect_files(str(tmp_path), [".xyz"])
        assert files == []


# ---------------------------------------------------------------------------
# BatchResult
# ---------------------------------------------------------------------------

class TestBatchResult:
    def test_summary_basic(self):
        r = BatchResult(total_files=10, processed_files=8,
                        pii_found_files=5, skipped_files=1, error_files=1)
        s = r.summary()
        assert "10" in s
        assert "8" in s
        assert "5" in s

    def test_summary_no_errors(self):
        r = BatchResult(total_files=3, processed_files=3, pii_found_files=1)
        s = r.summary()
        assert "錯誤" not in s


# ---------------------------------------------------------------------------
# run_batch
# ---------------------------------------------------------------------------

class TestRunBatch:
    def test_basic_batch(self, tmp_path):
        _create_test_tree(tmp_path)
        input_dir = str(tmp_path)
        output_dir = str(tmp_path / "output")

        config = {
            "custom_terms": {},
            "file_types": [".txt", ".md"],
            "substring_match": True,
        }

        result = run_batch(input_dir, output_dir, config,
                           reversible=True, use_ner=False)

        assert result.total_files == 3
        assert result.processed_files == 3
        assert os.path.isdir(output_dir)
        # Subfolder structure preserved
        assert os.path.isfile(os.path.join(output_dir, "sub", "doc3.txt"))

    def test_default_output_dir(self, tmp_path):
        src = tmp_path / "mydata"
        src.mkdir()
        (src / "test.txt").write_text("test 0912345678", encoding="utf-8")

        config = {"custom_terms": {}, "file_types": [".txt"], "substring_match": True}
        result = run_batch(str(src), None, config, use_ner=False)

        expected_out = str(src) + "_anonymized"
        assert os.path.isdir(expected_out)

    def test_progress_callback(self, tmp_path):
        (tmp_path / "a.txt").write_text("hello", encoding="utf-8")
        (tmp_path / "b.txt").write_text("world", encoding="utf-8")

        config = {"custom_terms": {}, "file_types": [".txt"], "substring_match": True}
        progress_log = []

        def on_progress(idx, total, fname):
            progress_log.append((idx, total, fname))

        run_batch(str(tmp_path), str(tmp_path / "out"), config,
                  use_ner=False, progress_callback=on_progress)

        assert len(progress_log) == 2
        assert progress_log[0][1] == 2  # total count

    def test_nonexistent_input_raises(self, tmp_path):
        with pytest.raises(ValueError, match="不存在"):
            run_batch("/nonexistent/dir", None, {}, use_ner=False)

    def test_pii_detection_in_batch(self, tmp_path):
        (tmp_path / "pii.txt").write_text("電話 0912345678", encoding="utf-8")

        config = {"custom_terms": {}, "file_types": [".txt"], "substring_match": True}
        result = run_batch(str(tmp_path), str(tmp_path / "out"), config, use_ner=False)

        assert result.pii_found_files >= 1
        # Output file should have been anonymized
        out_content = open(os.path.join(str(tmp_path), "out", "pii.txt"), encoding="utf-8").read()
        assert "0912345678" not in out_content

    def test_docx_output_remains_valid_docx(self, tmp_path):
        from docx import Document

        src = tmp_path / "sample.docx"
        doc = Document()
        doc.add_paragraph("電話 0912345678")
        doc.save(src)

        config = {"custom_terms": {}, "file_types": [".docx"], "substring_match": True}
        result = run_batch(str(tmp_path), str(tmp_path / "out"), config, use_ner=False)

        assert result.pii_found_files == 1
        out_path = tmp_path / "out" / "sample.docx"
        out_doc = Document(str(out_path))
        text = "\n".join(p.text for p in out_doc.paragraphs)
        assert "0912345678" not in text
        assert "__ANON:PHONE_001__" in text

    def test_empty_folder(self, tmp_path):
        empty = tmp_path / "empty"
        empty.mkdir()

        config = {"custom_terms": {}, "file_types": [".txt"], "substring_match": True}
        result = run_batch(str(empty), None, config, use_ner=False)

        assert result.total_files == 0
        assert result.processed_files == 0
