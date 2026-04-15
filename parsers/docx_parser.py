from docx import Document


class DocxParser:
    EXTENSIONS = {'.docx'}
    OUTPUT_EXTENSION = ".docx"

    def parse(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
        except Exception:
            return ""
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return '\n'.join(parts)

    def anonymize_to_path(self, file_path: str, output_path: str, anonymizer):
        try:
            doc = Document(file_path)
        except Exception:
            return False, []

        spans = []

        def anonymize_paragraphs(paragraphs):
            for para in paragraphs:
                anonymized, para_spans = anonymizer._anonymize_text_with_spans(para.text)
                if para_spans:
                    para.text = anonymized
                    spans.extend(para_spans)

        def anonymize_tables(tables):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        anonymized, cell_spans = anonymizer._anonymize_text_with_spans(cell.text)
                        if cell_spans:
                            cell.text = anonymized
                            spans.extend(cell_spans)

        anonymize_paragraphs(doc.paragraphs)
        anonymize_tables(doc.tables)

        for section in doc.sections:
            anonymize_paragraphs(section.header.paragraphs)
            anonymize_tables(section.header.tables)
            anonymize_paragraphs(section.footer.paragraphs)
            anonymize_tables(section.footer.tables)

        if not spans:
            return False, []

        doc.save(output_path)
        return True, spans
